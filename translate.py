from docx import Document
import os

# Путь к папке "Загрузки" на Windows
downloads_path = os.path.join(os.path.expanduser('~'), 'Downloads')

# Создаем новый документ
doc = Document()

# Заголовок документа
doc.add_heading('ОТДЕЛ КОНТРОЛЯ КАЧЕСТВА', level=1)
doc.add_heading('СЕРТИФИКАТ АНАЛИЗА', level=1)
doc.add_heading('ГОТОВАЯ ПРОДУКЦИЯ', level=1)

# Информация о продукте
doc.add_paragraph('Наименование продукта: ПРОПРАНОЛОЛА ГИДРОХЛОРИД, ВЕЩЕСТВО-ПОРОШОК')
doc.add_paragraph('Номер партии: 24040PRRII')
doc.add_paragraph('Размер партии: 537,70 кг')
doc.add_paragraph('Дата производства: Июнь/2024')
doc.add_paragraph('Дата повторного тестирования: Нет')
doc.add_paragraph('Срок годности: Май/2029')
doc.add_paragraph('Номер анализа: IBDA-242249')
doc.add_paragraph('Номер LIMS: 2102-44398')
doc.add_paragraph('Идентификатор спецификации LIMS: 2102-639')
doc.add_paragraph('Страница: 1 из 3')
doc.add_paragraph('Номер спецификации/Редакция: TS/API/PR/RUS/3.0/0')

# Таблица с результатами анализа
table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Тесты'
hdr_cells[1].text = 'Спецификации'
hdr_cells[2].text = 'Результаты'

# Добавление данных в таблицу
data = [
    ("Описание", "Белый или почти белый мелкокристаллический порошок.", "Белый мелкокристаллический порошок"),
    ("Растворимость", "Растворим в воде и в этаноле 96%, практически нерастворим в гептане.", "Соответствует"),
    ("Идентификация 1", "ИК-спектрометрия: ИК-спектр вещества должен соответствовать ИК-спектру стандартного образца пропранолола гидрохлорида.", "Соответствует"),
    ("Идентификация 2", "Качественная реакция на хлориды: Соответствует требованиям.", "Соответствует"),
    ("Температура плавления", "От 163°C до 166°C.", "165°C"),
    ("Удельное вращение", "От -1,0 до +1,0 (4% раствор вещества в воде).", "+0,0"),
    ("Прозрачность и цвет раствора", "Прозрачность раствора: 10% раствор вещества в метаноле должен быть прозрачным.", "Соответствует")
]

for test, spec, result in data:
    row_cells = table.add_row().cells
    row_cells[0].text = test
    row_cells[1].text = spec
    row_cells[2].text = result

# Подписи и даты
doc.add_paragraph('Проверено (ОК): АМИТ СИНГХ (Менеджер ОК)')
doc.add_paragraph('Утверждено (ОА): ШАЛИНИ БХАНГУ (Менеджер ОА)')
doc.add_paragraph('Проверено: 25/07/2024 17:14')
doc.add_paragraph('Утверждено: 26/07/2024 16:51')
doc.add_paragraph('Напечатано: ШАЛИНИ БХАНГУ')
doc.add_paragraph('Напечатано: 26/07/2024 16:51')
doc.add_paragraph('Копия №: 1')

# Примечание
doc.add_paragraph('Примечание: Этот документ был сгенерирован электронным способом и действителен без подписи.')

# Контактная информация
doc.add_paragraph('Ipca Laboratories Ltd.')
doc.add_paragraph('P.O. Sejavta, RATLAM 457 001, Индия | Тел.: +91 07412-60929/60445/50 Факс: +91 0741207412-60783')
doc.add_paragraph('Зарегистрированный офис: 48, Kandivli Industrial Estate, Kandivli (West), Mumbai 400067, Индия | Тел.: +91 22 6647 4444 Факс: +91 22 2868 6613')
doc.add_paragraph('E-mail: ipca@ipca.com | CIN: L24239MH1949PLC007837')

# Сохраняем документ в папку "Загрузки"
file_path = os.path.join(downloads_path, 'Сертификат_анализа_пропранолол.docx')
doc.save(file_path)

print(f"Документ сохранен как '{file_path}'")